import re
from pathlib import Path
import pdfplumber
from docx import Document

SKILLS = ["python", "java", "javascript", "typescript", "react", "node.js", "fastapi", "django", "flask", "sql", "sqlite", "mysql", "postgresql", "mongodb", "aws", "docker", "kubernetes", "git", "github", "selenium", "pytest", "postman", "jira", "manual testing", "api testing", "pandas", "numpy", "excel", "power bi", "tableau", "scikit-learn", "machine learning", "html", "css", "tailwind", "redux", "rest api", "linux"]
SECTION_NAMES = ["education", "experience", "projects", "certifications"]

def extract_text(path: Path) -> str:
    if path.suffix.lower() == ".pdf":
        with pdfplumber.open(path) as pdf:
            return "\n".join(page.extract_text() or "" for page in pdf.pages)
    doc = Document(path)
    return "\n".join(p.text for p in doc.paragraphs) + "\n" + "\n".join(" ".join(cell.text for cell in row.cells) for t in doc.tables for row in t.rows)

def _section(text: str, header: str) -> list[str]:
    pattern = rf"(?ims)^\s*{re.escape(header)}\s*$\s*(.*?)(?=^\s*(?:{'|'.join(SECTION_NAMES)})\s*$|\Z)"
    found = re.search(pattern, text)
    if not found: return []
    return [line.strip(" •-\t") for line in found.group(1).splitlines() if line.strip()]

def parse_resume(text: str) -> dict:
    lines = [x.strip() for x in text.splitlines() if x.strip()]
    email = re.search(r"[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}", text)
    phone = re.search(r"(?:\+?\d{1,3}[ .-]?)?(?:\(?\d{2,4}\)?[ .-]?)?\d{3,5}[ .-]?\d{4}", text)
    lower = text.lower()
    skills = [skill for skill in SKILLS if re.search(rf"(?<!\w){re.escape(skill)}(?!\w)", lower)]
    name = next((x for x in lines[:5] if not email or email.group(0) not in x and len(x) < 60 and not re.search(r"\d", x)), "Not detected")
    return {"name": name, "email": email.group(0) if email else None, "phone": phone.group(0) if phone else None, "skills": skills, "education": _section(text, "education"), "experience": _section(text, "experience") + _section(text, "work experience"), "projects": _section(text, "projects"), "certifications": _section(text, "certifications")}
